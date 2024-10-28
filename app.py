from flask import Flask, render_template, request, send_file, jsonify
import requests
from io import BytesIO
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

app = Flask(__name__)

# Fungsi untuk memuat data dari URL
def load_data():
    response = requests.get("https://stamet-juanda.bmkg.go.id/cuwis/json/test.json")
    return response.json()

@app.route('/')
def index():
    data = load_data()
    
    # Mengambil daftar kota dan kecamatan yang unik
    kota_list = sorted(set(item['Kota'] for item in data))
    
    # Mengambil kecamatan per kota
    kecamatan_dict = {}
    for item in data:
        if item['Kota'] not in kecamatan_dict:
            kecamatan_dict[item['Kota']] = set()
        kecamatan_dict[item['Kota']].add(item['Kecamatan'])
    
    return render_template('index.html', kota_list=kota_list, kecamatan_dict=kecamatan_dict)

@app.route('/search', methods=['POST'])
def search():
    data = load_data()
    tanggal = request.form['tanggal']
    kota = request.form['kota']
    kecamatan = request.form['kecamatan']
    
    # Mencari data berdasarkan tanggal, kota, dan kecamatan
    result = next((item for item in data if item['Tanggal'] == tanggal and item['Kota'] == kota and item['Kecamatan'] == kecamatan), None)
    
    return render_template('index.html', kota_list=sorted(set(item['Kota'] for item in data)),
                           kecamatan_dict=sorted(set(item['Kecamatan'] for item in data)),
                           result=result)

@app.route('/get_kecamatan/<kota>', methods=['GET'])
def get_kecamatan(kota):
    data = load_data()
    kecamatan_list = sorted(set(item['Kecamatan'] for item in data if item['Kota'] == kota))
    return jsonify(kecamatan_list)

@app.route('/download/<file_type>', methods=['POST'])
def download(file_type):
    data = load_data()
    tanggal = request.form['tanggal']
    kota = request.form['kota']
    kecamatan = request.form['kecamatan']

    # Mencari data berdasarkan tanggal, kota, dan kecamatan
    result = next((item for item in data if item['Tanggal'] == tanggal and item['Kota'] == kota and item['Kecamatan'] == kecamatan), None)
    
    if result:
        if file_type == 'word':
            # Membuat dokumen Word
            doc = Document()
            doc.add_heading(f'Laporan Cuaca: {result["Tanggal"]} - {result["Kota"]}, {result["Kecamatan"]}', level=1)
            doc.add_paragraph(f'Tanggal: {result["Tanggal"]}')
            doc.add_paragraph(f'ID: 5008350')  # Sesuaikan dengan ID yang relevan
            doc.add_paragraph(f'Kota: {result["Kota"]}')
            doc.add_paragraph(f'Kecamatan: {result["Kecamatan"]}')
            doc.add_paragraph(f'Cuaca Pagi: {result["Cuaca_Pagi"]}')
            doc.add_paragraph(f'Cuaca Siang: {result["Cuaca_Siang"]}')
            doc.add_paragraph(f'Cuaca Malam: {result["Cuaca_Malam"]}')
            doc.add_paragraph(f'Cuaca Dini: {result["Cuaca_Dini"]}')
            doc.add_paragraph(f'Suhu Min: {result["Suhu_Min"]} °C')
            doc.add_paragraph(f'Suhu Max: {result["Suhu_Max"]} °C')
            doc.add_paragraph(f'Kelembapan Min: {result["Kelembapan_Min"]} %')
            doc.add_paragraph(f'Kelembapan Max: {result["Kelembapan_Max"]} %')
            doc.add_paragraph(f'Arah Angin: {result["Arah_Angin"]}')
            doc.add_paragraph(f'Kecepatan Angin Min: {result["Kecepatan_Angin_Min"]} km/h')
            doc.add_paragraph(f'Kecepatan Angin Max: {result["Kecepatan_Angin_Max"]} km/h')

            # Simpan dokumen ke dalam buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            return send_file(buffer, as_attachment=True, download_name='Laporan_Cuaca.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

        elif file_type == 'pdf':
            # Membuat dokumen PDF
            buffer = BytesIO()
            c = canvas.Canvas(buffer, pagesize=letter)
            c.drawString(100, 750, f'Laporan Cuaca: {result["Tanggal"]} - {result["Kota"]}, {result["Kecamatan"]}')
            c.drawString(100, 735, f'Tanggal: {result["Tanggal"]}')
            c.drawString(100, 720, f'ID: 5008350')  # Sesuaikan dengan ID yang relevan
            c.drawString(100, 705, f'Kota: {result["Kota"]}')
            c.drawString(100, 690, f'Kecamatan: {result["Kecamatan"]}')
            c.drawString(100, 675, f'Cuaca Pagi: {result["Cuaca_Pagi"]}')
            c.drawString(100, 660, f'Cuaca Siang: {result["Cuaca_Siang"]}')
            c.drawString(100, 645, f'Cuaca Malam: {result["Cuaca_Malam"]}')
            c.drawString(100, 630, f'Cuaca Dini: {result["Cuaca_Dini"]}')
            c.drawString(100, 615, f'Suhu Min: {result["Suhu_Min"]} °C')
            c.drawString(100, 600, f'Suhu Max: {result["Suhu_Max"]} °C')
            c.drawString(100, 585, f'Kelembapan Min: {result["Kelembapan_Min"]} %')
            c.drawString(100, 570, f'Kelembapan Max: {result["Kelembapan_Max"]} %')
            c.drawString(100, 555, f'Arah Angin: {result["Arah_Angin"]}')
            c.drawString(100, 540, f'Kecepatan Angin Min: {result["Kecepatan_Angin_Min"]} km/h')
            c.drawString(100, 525, f'Kecepatan Angin Max: {result["Kecepatan_Angin_Max"]} km/h')
            c.showPage()
            c.save()
            buffer.seek(0)

            return send_file(buffer, as_attachment=True, download_name='Laporan_Cuaca.pdf', mimetype='application/pdf')

    return "Data tidak ditemukan", 404

if __name__ == '__main__':
    app.run(debug=True)